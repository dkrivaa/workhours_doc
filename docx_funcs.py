import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date
from io import BytesIO


from helpers import months, month_names_dict


def reorder_dataframe(df):
    # Reset index
    df = df.reset_index(drop=True)
    # Change order of columns
    column_order = ['total', 'End Hour', 'Start Hour', 'Comments', 'Date']
    df = df[column_order]
    # Rename columns to hebrew
    df.columns = ['סך הכל שעות', 'שעת סיום', 'שעת התחלה', '', 'תאריך']
    return df


def calc_total_hours(df):
    # Calculate total hours without modifying the DataFrame
    total_time = pd.to_timedelta(df['total'] + ':00').sum()  # Adds ":00" for seconds if necessary
    total_hours = total_time.total_seconds() / 3600  # Convert total seconds to hours
    return total_hours


def set_rtl(paragraph):
    """Apply Right-to-Left formatting to a given paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def add_paragraph_with_text(document, text, style=None, alignment=None, space_after=None, rtl=False):
    """Add a styled and aligned paragraph with optional RTL formatting."""
    paragraph = document.add_paragraph(text)
    if style:
        paragraph.style = style
    if alignment:
        paragraph.alignment = alignment
    if space_after:
        paragraph.paragraph_format.space_after = Pt(space_after)
    if rtl:
        set_rtl(paragraph)
    return paragraph


def create_document(df, title_text, total_hours_text):
    document = Document()

    # Set document-level Right-to-Left (RTL) layout
    section = document.sections[0]
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    section._sectPr.append(bidi)

    # Today's date in format dd/mm/yyyy
    today = date.today().strftime("%d/%m/%Y")

    # Add date paragraph with right alignment
    add_paragraph_with_text(document, today, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=20)

    # Add subject title in center with Heading1 style and RTL
    add_paragraph_with_text(document, title_text, style='Heading1', alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30,
                            rtl=True)

    # Add table with header row and data rows
    table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = 'Table Grid'

    # Populate header row
    for j, column_name in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(column_name)
        set_rtl(cell.paragraphs[0])

    # Populate data rows
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = str(value)
            set_rtl(cell.paragraphs[0])

    # Add a blank paragraph after the table
    document.add_paragraph()

    # Add total hours paragraph with left alignment
    add_paragraph_with_text(document, total_hours_text, style='Heading1', alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            space_after=20, rtl=True)

    # Add blessing paragraph in center with Heading2 style and RTL
    add_paragraph_with_text(document, 'בברכה', style='Heading2', alignment=WD_ALIGN_PARAGRAPH.CENTER, rtl=True)

    # Save document to a BytesIO buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer


def one_docx(df):

    title_text = ''
    month_list, year_list = months(df)
    month_dict = month_names_dict()

    if len(month_list) == 1:
        month_name = month_dict[str(month_list[0])]
        year = year_list[0]
        st.write(month_name)
        hebrew_text = 'דיווח שעות לחודש'
        title_text = f'{hebrew_text} {month_name} {year} '

    elif len(month_list) > 1:
        min_month = month_dict[str(month_list[0])]
        max_month = month_dict[str(month_list[-1])]
        hebrew_text = 'דיווח שעות לחודשים'
        if len(list(set(year_list))) == 1:
            year = year_list[0]
            title_text = f'{hebrew_text} {min_month}-{max_month} {year} '
        elif len(list(set(year_list))) > 1:
            title_text = f'{hebrew_text} {min_month}-{max_month} {year_list[0]}-{year_list[-1]} '

    total_hours = calc_total_hours(df)
    hebrew_total = 'סך שעות בדוח'
    total_hours_text = f'{hebrew_total} {total_hours} :'

    df = reorder_dataframe(df)

    buffer = create_document(df, title_text, total_hours_text)

    return buffer



